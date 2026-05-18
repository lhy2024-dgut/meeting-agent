from pathlib import Path

from docx import Document
from docxtpl import DocxTemplate, RichText

import config
from engines.pdf_engine import PDFEngine


class ExportChain:
    """文档导出链，支持 docx / md / pdf"""

    def __init__(self):
        self.output_dir = config.OUTPUT_DIR
        self.pdf_engine = PDFEngine()

    def run(self, data, output_format="docx", template_path=None):
        context = {
            "title": data.get("title", "会议纪要"),
            "date": data.get("date", ""),
            "minutes_content": data.get("minutes", ""),
            "action_items": data.get("action_items", ""),
            "resolutions": data.get("resolutions", ""),
        }
        output_path = self.output_dir / f"meeting_{data.get('meeting_id', 'output')}"
        fmt = (output_format or "docx").lower()

        if fmt == "docx":
            output_path = output_path.with_suffix(".docx")
            if template_path and template_path.lower().endswith(".docx"):
                return self._fill_docx_template(template_path, context, output_path)
            return self._create_docx(context["minutes_content"], output_path)

        if fmt == "md":
            output_path = output_path.with_suffix(".md")
            return self._create_md(context["minutes_content"], output_path)

        if fmt == "pdf":
            output_path = output_path.with_suffix(".pdf")
            if template_path and template_path.lower().endswith(".pdf"):
                return self.pdf_engine.fill_template_with_context(
                    template_path, context, str(output_path)
                )
            return self.pdf_engine.create_pdf_from_markdown(
                context["minutes_content"], str(output_path)
            )

        raise ValueError(f"不支持的输出格式: {fmt}")

    def _fill_docx_template(self, template_path, context, output_path):
        doc = DocxTemplate(template_path)
        rich_context = {}
        for key, value in context.items():
            if isinstance(value, str) and key in {
                "minutes_content", "action_items", "resolutions"
            }:
                rt = RichText()
                for line in value.split("\n"):
                    if line.startswith("# "):
                        rt.add(line[2:] + "\n", bold=True, size=16)
                    elif line.startswith("## "):
                        rt.add(line[3:] + "\n", bold=True, size=14)
                    elif line.startswith("- "):
                        rt.add("• " + line[2:] + "\n")
                    else:
                        rt.add(line + "\n")
                rich_context[key] = rt
            else:
                rich_context[key] = value
        doc.render(rich_context)
        doc.save(str(output_path))
        return str(output_path)

    def _create_docx(self, md_content, output_path):
        doc = Document()
        for line in (md_content or "").split("\n"):
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            else:
                doc.add_paragraph(line)
        doc.save(str(output_path))
        return str(output_path)

    def _create_md(self, content, output_path):
        output_path.write_text(content or "", encoding="utf-8")
        return str(output_path)
