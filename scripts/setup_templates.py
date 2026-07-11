"""任务 3.2 — 生成 Word/PDF 模板文件 + 预览图（仅首次运行）"""

import subprocess
import sys
from pathlib import Path

# 确保依赖已安装
for pkg in ["docxtpl", "Pillow"]:
    try:
        __import__(pkg.replace("Pillow", "PIL"))
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", pkg])

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from PIL import Image, ImageDraw, ImageFont

BASE = Path("storage/templates/source")
PREVIEW_DIR = Path("storage/templates/previews")
BASE.mkdir(parents=True, exist_ok=True)
PREVIEW_DIR.mkdir(parents=True, exist_ok=True)


# =====================================================================
#  样式工具
# =====================================================================

def _set_run_font(run, name_cn="微软雅黑", size=11, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name_cn
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), name_cn)
    if color:
        run.font.color.rgb = color


def _add_cover(doc, title="会议纪要", subtitle=""):
    """添加封面"""
    for _ in range(6):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title or "会议纪要")
    _set_run_font(run, "微软雅黑", 28, bold=True, color=RGBColor(0x1A, 0x1A, 0x2E))

    if subtitle:
        doc.add_paragraph("")
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        _set_run_font(r2, "微软雅黑", 14, color=RGBColor(0x64, 0x74, 0x8B))

    doc.add_page_break()


def _add_heading_styled(doc, text, level=1):
    """添加带格式的标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        _set_run_font(run, "微软雅黑", 18, bold=True, color=RGBColor(0x1A, 0x1A, 0x2E))
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
    elif level == 2:
        _set_run_font(run, "微软雅黑", 14, bold=True, color=RGBColor(0x2D, 0x3A, 0x5E))
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
    else:
        _set_run_font(run, "微软雅黑", 12, bold=True, color=RGBColor(0x47, 0x58, 0x7A))


def _add_para(doc, text, size=11, bold=False, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    _set_run_font(run, "微软雅黑", size, bold=bold, color=color)
    return p


def _add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("—" * 40)
    _set_run_font(run, "微软雅黑", 8, color=RGBColor(0xCC, 0xCC, 0xCC))


def _set_page_margins(doc, top=2.5, bottom=2.5, left=2.8, right=2.8):
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def _create_base_doc(title="会议纪要", subtitle="会议日期：{{date}}"):
    doc = Document()
    _set_page_margins(doc)

    # 封面
    _add_cover(doc, title, subtitle)

    # 目录占位
    _add_heading_styled(doc, "目  录", level=1)
    _add_para(doc, "（在 Word 中右键此处 → 更新域 可自动生成目录）", size=9, color=RGBColor(0x99, 0x99, 0x99))
    _add_para(doc, "请在 Word 中插入自动目录：引用 → 目录 → 自动目录。", size=9, color=RGBColor(0x99, 0x99, 0x99))
    doc.add_page_break()

    return doc


# =====================================================================
#  DEFAULT 模板 — 商务稳重
# =====================================================================

def create_default_docx():
    doc = _create_base_doc("{{title}}", "会议日期：{{date}}")

    # 1. 纪要正文
    _add_heading_styled(doc, "一、会议纪要", level=1)
    _add_divider(doc)
    _add_para(doc, "{{minutes_content}}")

    doc.add_page_break()

    # 2. 待办事项
    _add_heading_styled(doc, "二、待办事项", level=1)
    _add_divider(doc)
    _add_para(doc, "{{action_items}}")

    doc.add_page_break()

    # 3. 决议
    _add_heading_styled(doc, "三、会议决议", level=1)
    _add_divider(doc)
    _add_para(doc, "{{resolutions}}")

    # 页脚：页码
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("— 第  页 —")
        _set_run_font(run, "微软雅黑", 9, color=RGBColor(0xAA, 0xAA, 0xAA))

    path = BASE / "default_meeting.docx"
    doc.save(str(path))
    print(f"[OK] 创建 {path}")
    return path


def create_elegant_docx():
    """ELEGANT 模板 — 简约清晰，左侧竖线装饰"""
    doc = _create_base_doc("{{title}}", "{{date}}")

    # 纪要正文
    _add_heading_styled(doc, "📝 会议记录", level=1)
    _add_divider(doc)
    _add_para(doc, "{{minutes_content}}")

    doc.add_page_break()

    # 待办事项 — 使用表格版式
    _add_heading_styled(doc, "✅ 待办事项", level=1)
    _add_divider(doc)
    _add_para(doc, "{{action_items}}")

    doc.add_page_break()

    # 决议
    _add_heading_styled(doc, "📌 决议", level=1)
    _add_divider(doc)
    _add_para(doc, "{{resolutions}}")

    # 页脚
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("会议纪要 · 自动生成")
        _set_run_font(run, "微软雅黑", 9, color=RGBColor(0xAA, 0xAA, 0xAA))

    path = BASE / "elegant_meeting.docx"
    doc.save(str(path))
    print(f"[OK] 创建 {path}")
    return path


# =====================================================================
#  PDF 模板（PyMuPDF）
# =====================================================================

def _get_font_path():
    """查找系统中文字体"""
    candidates = [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/NotoSansSC-VF.ttf",
    ]
    for p in candidates:
        if Path(p).exists():
            return p
    return None


def create_default_pdf():
    import fitz
    font_path = _get_font_path()

    doc = fitz.open()
    page_w, page_h = 595, 842  # A4

    # 封面
    page = doc.new_page(width=page_w, height=page_h)
    rect = fitz.Rect(50, 250, 545, 400)
    page.insert_textbox(
        rect, "{{title}}", fontsize=28,
        fontfile=font_path,
        color=(0.1, 0.1, 0.18), align=1,
    )
    rect2 = fitz.Rect(50, 410, 545, 450)
    page.insert_textbox(
        rect2, "{{date}}", fontsize=14,
        fontfile=font_path,
        color=(0.39, 0.45, 0.55), align=1,
    )

    # 内容页
    page2 = doc.new_page(width=page_w, height=page_h)
    # 标题
    page2.insert_textbox(
        fitz.Rect(50, 50, 545, 100), "会议纪要", fontsize=22,
        fontfile=font_path,
        color=(0.1, 0.1, 0.18),
    )
    # 分隔线
    page2.draw_line(fitz.Point(50, 105), fitz.Point(545, 105), color=(0.8, 0.8, 0.8), width=1)
    # 正文
    page2.insert_textbox(
        fitz.Rect(50, 120, 545, 800), "{{minutes_content}}", fontsize=11,
        fontfile=font_path,
        color=(0.2, 0.2, 0.2),
    )

    # 待办页
    page3 = doc.new_page(width=page_w, height=page_h)
    page3.insert_textbox(
        fitz.Rect(50, 50, 545, 100), "待办事项", fontsize=22,
        fontfile=font_path,
        color=(0.1, 0.1, 0.18),
    )
    page3.draw_line(fitz.Point(50, 105), fitz.Point(545, 105), color=(0.8, 0.8, 0.8), width=1)
    page3.insert_textbox(
        fitz.Rect(50, 120, 545, 800), "{{action_items}}", fontsize=11,
        fontfile=font_path,
        color=(0.2, 0.2, 0.2),
    )

    # 决议页
    page4 = doc.new_page(width=page_w, height=page_h)
    page4.insert_textbox(
        fitz.Rect(50, 50, 545, 100), "会议决议", fontsize=22,
        fontfile=font_path,
        color=(0.1, 0.1, 0.18),
    )
    page4.draw_line(fitz.Point(50, 105), fitz.Point(545, 105), color=(0.8, 0.8, 0.8), width=1)
    page4.insert_textbox(
        fitz.Rect(50, 120, 545, 800), "{{resolutions}}", fontsize=11,
        fontfile=font_path,
        color=(0.2, 0.2, 0.2),
    )

    path = BASE / "default_meeting.pdf"
    doc.save(str(path))
    doc.close()
    print(f"[OK] 创建 {path}")
    return path


def create_elegant_pdf():
    import fitz
    font_path = _get_font_path()

    doc = fitz.open()
    page_w, page_h = 595, 842

    # 封面 — 带装饰线
    page = doc.new_page(width=page_w, height=page_h)
    # 左侧装饰线
    page.draw_line(fitz.Point(50, 200), fitz.Point(50, 500), color=(0.2, 0.5, 0.8), width=3)
    page.insert_textbox(
        fitz.Rect(70, 250, 545, 380), "{{title}}", fontsize=26,
        fontfile=font_path,
        color=(0.1, 0.1, 0.18),
    )
    page.insert_textbox(
        fitz.Rect(70, 390, 545, 430), "{{date}}", fontsize=13,
        fontfile=font_path,
        color=(0.39, 0.45, 0.55),
    )

    # 内容页
    page2 = doc.new_page(width=page_w, height=page_h)
    page2.insert_textbox(
        fitz.Rect(50, 50, 545, 95), "📝 会议记录", fontsize=20,
        fontfile=font_path,
        color=(0.1, 0.1, 0.18),
    )
    page2.draw_line(fitz.Point(50, 100), fitz.Point(350, 100), color=(0.2, 0.5, 0.8), width=2)
    page2.insert_textbox(
        fitz.Rect(50, 115, 545, 800), "{{minutes_content}}", fontsize=11,
        fontfile=font_path,
        color=(0.2, 0.2, 0.2),
    )

    path = BASE / "elegant_meeting.pdf"
    doc.save(str(path))
    doc.close()
    print("[OK] 创建", path)
    return path


# =====================================================================
#预览图
# =====================================================================

def create_preview(name, title_text="会议纪要模板"):
    img = Image.new("RGB", (480, 360), (255, 255, 255))
    draw = ImageDraw.Draw(img)

    # 尝试加载字体
    try:
        font_large = ImageFont.truetype("C:/Windows/Fonts/msyh.ttc", 22)
        font_small = ImageFont.truetype("C:/Windows/Fonts/msyh.ttc", 13)
    except Exception:
        font_large = ImageFont.load_default()
        font_small = ImageFont.load_default()

    # 封面块
    draw.rectangle([20, 20, 460, 140], fill=(245, 247, 250), outline=(200, 210, 220), width=1)
    draw.text((240, 50), title_text, fill=(30, 30, 50), font=font_large, anchor="mt")
    draw.text((240, 90), "会议标题 / 日期", fill=(100, 110, 130), font=font_small, anchor="mt")

    # 内容块
    colors = [(235, 240, 255), (245, 245, 250), (240, 250, 240)]
    labels = ["[纪要正文]", "[待办事项]", "[决议]"]
    for i in range(3):
        y = 160 + i * 65
        draw.rectangle([20, y, 460, y + 55], fill=colors[i], outline=(200, 210, 220), width=1)
        draw.text((40, y + 18), labels[i], fill=(50, 50, 70), font=font_small, anchor="lt")
        draw.text((200, y + 18), "━━━━━━━━━━━━━━━━", fill=(180, 185, 195), font=font_small, anchor="lt")

    # 页脚
    draw.line([(20, 345), (460, 345)], fill=(220, 220, 220), width=1)
    draw.text((240, 350), "会议纪要 - 自动生成", fill=(170, 170, 170), font=font_small, anchor="mt")

    path = PREVIEW_DIR / f"{name}.png"
    img.save(str(path))
    print(f"[OK] 创建 {path}")


# =====================================================================
#  主流程
# =====================================================================

if __name__ == "__main__":
    print("=" * 50)
    print("[模板生成]")
    print("=" * 50)

    create_default_docx()
    create_elegant_docx()
    create_default_pdf()
    create_elegant_pdf()

    create_preview("default", "默认模板 - 商务稳重")
    create_preview("elegant", "优雅模板 - 简约清晰")

    print("=" * 50)
    print("[全部模板文件已生成]")
    print(f"  [dir] {BASE.resolve()}")
    print(f"  [dir] {PREVIEW_DIR.resolve()}")
