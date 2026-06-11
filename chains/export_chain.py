from pathlib import Path

from docx import Document
from docxtpl import DocxTemplate, RichText

import config
from chains.minutes_chain import PLACEHOLDER_ALL_EMPTY
from engines.pdf_engine import PDFEngine
from logger import get_logger

logger = get_logger(__name__)

# ── 内置模板路径 ──────────────────────────────────────────────
TEMPLATE_SOURCE_DIR = Path("storage/templates/source")
TEMPLATE_PREVIEW_DIR = Path("storage/templates/previews")

_TEMPLATE_REGISTRY = {}  # name -> {"docx": Path, "pdf": Path, "preview": Path}


def _scan_templates():
    """扫描 storage/templates/source/ 下的模板文件"""
    _TEMPLATE_REGISTRY.clear()
    if not TEMPLATE_SOURCE_DIR.exists():
        return

    docx_files = {}
    pdf_files = {}
    for f in TEMPLATE_SOURCE_DIR.iterdir():
        if f.suffix == ".docx":
            stem = f.stem.replace("_meeting", "").replace(".docx", "")
            docx_files[stem] = f
        elif f.suffix == ".pdf":
            stem = f.stem.replace("_meeting", "").replace(".pdf", "")
            pdf_files[stem] = f

    all_names = set(docx_files.keys()) | set(pdf_files.keys())
    for name in sorted(all_names):
        _TEMPLATE_REGISTRY[name] = {
            "docx": docx_files.get(name),
            "pdf": pdf_files.get(name),
            "preview": TEMPLATE_PREVIEW_DIR / f"{name}.png",
        }


def list_templates() -> list[dict]:
    """返回可用模板列表，每个模板含 name / label / preview_path"""
    _scan_templates()
    labels = {
        "default": "默认模板 - 商务稳重",
        "elegant": "优雅模板 - 简约清晰",
    }
    result = []
    for name, paths in _TEMPLATE_REGISTRY.items():
        preview = paths["preview"] if paths["preview"].exists() else None
        result.append({
            "name": name,
            "label": labels.get(name, name),
            "has_docx": paths["docx"] is not None,
            "has_pdf": paths["pdf"] is not None,
            "preview_path": str(preview) if preview else None,
        })
    return result


def get_template_path(name: str, fmt: str) -> str | None:
    """获取指定模板的路径，fmt='docx' 或 'pdf'，不存在返回 None"""
    _scan_templates()
    paths = _TEMPLATE_REGISTRY.get(name)
    if not paths:
        return None
    p = paths.get(fmt)
    return str(p) if p and p.exists() else None


class ExportChain:
    """文档导出链，支持 docx / md / pdf"""

    def __init__(self, output_dir=None, pdf_engine=None):
        self.output_dir = output_dir or config.OUTPUT_DIR
        self.pdf_engine = pdf_engine or PDFEngine()

    def run(self, data, output_format="docx", template_path=None, template_name=None):
        context = {
            "title": data.get("title", "会议纪要"),
            "date": data.get("date", ""),
            "minutes_content": data.get("minutes", ""),
            "action_items": data.get("action_items", ""),
            "resolutions": data.get("resolutions", ""),
        }
        output_path = self.output_dir / f"meeting_{data.get('meeting_id', 'output')}"
        fmt = (output_format or "docx").lower()

        # ── 若指定了 template_name 且未传 template_path，自动查找内置模板 ──
        if template_name and not template_path:
            template_path = get_template_path(template_name, fmt)

        if fmt == "docx":
            output_path = output_path.with_suffix(".docx")
            if template_path and template_path.lower().endswith(".docx"):
                return self._fill_docx_template(template_path, context, output_path)
            return self._create_docx(context["minutes_content"], output_path)

        if fmt == "md":
            output_path = output_path.with_suffix(".md")
            return self._create_md(context, output_path)

        if fmt == "pdf":
            output_path = output_path.with_suffix(".pdf")
            # 构建完整的 Markdown 内容（标题 + 正文 + 待办 + 决议）
            full_md = self._build_full_markdown(context)
            if template_path and template_path.lower().endswith(".pdf"):
                return self.pdf_engine.fill_template_with_context(
                    template_path, context, str(output_path)
                )
            return self.pdf_engine.create_pdf_from_markdown(
                full_md, str(output_path)
            )

        raise ValueError(f"不支持的输出格式: {fmt}")

    def _fill_docx_template(self, template_path, context, output_path):
        doc = DocxTemplate(template_path)
        rich_context = {}
        for key, value in context.items():
            if isinstance(value, str) and key in {
                "minutes_content", "action_items", "resolutions"
            }:
                rt = RichText()
                for line in value.split("\n"):
                    if line.startswith("# "):
                        rt.add(line[2:] + "\n", bold=True, size=16)
                    elif line.startswith("## "):
                        rt.add(line[3:] + "\n", bold=True, size=14)
                    elif line.startswith("- "):
                        rt.add("• " + line[2:] + "\n")
                    else:
                        rt.add(line + "\n")
                rich_context[key] = rt
            else:
                rich_context[key] = value
        doc.render(rich_context)
        doc.save(str(output_path))
        return str(output_path)

    def _create_docx(self, md_content, output_path):
        doc = Document()
        for line in (md_content or "").split("\n"):
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            else:
                doc.add_paragraph(line)
        doc.save(str(output_path))
        return str(output_path)

    def _build_full_markdown(self, context: dict) -> str:
        """从 context 构建完整的 Markdown 文档"""
        lines = []
        title = context.get("title", "会议纪要")
        date = context.get("date", "")
        minutes_content = context.get("minutes_content", "")
        action_items = context.get("action_items", "")
        resolutions = context.get("resolutions", "")

        lines.append(f"# {title}")
        if date:
            lines.append(f"\n**日期**：{date}\n")
        lines.append(minutes_content or "")

        if action_items and action_items.strip() not in PLACEHOLDER_ALL_EMPTY:
            lines.append("\n---\n## 待办事项\n")
            lines.append(action_items)

        if resolutions and resolutions.strip() not in PLACEHOLDER_ALL_EMPTY:
            lines.append("\n---\n## 会议决议\n")
            lines.append(resolutions)

        return "\n".join(lines).strip()

    def _create_md(self, context, output_path):
        content = self._build_full_markdown(context)
        output_path.write_text(content or "", encoding="utf-8")
        return str(output_path)
