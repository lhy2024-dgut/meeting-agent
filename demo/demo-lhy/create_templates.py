"""
create_templates.py - 生成初始模板文件
运行一次即可，生成 templates/ 目录下的所有模板
"""

import os

os.makedirs("templates", exist_ok=True)

# ── 生成 Word 模板
def create_word_template():
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{ title }}")
    run.bold = True
    run.font.size = Pt(20)

    # 日期
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("日期：{{ date }}")

    doc.add_paragraph()

    sections = [
        ("会议纪要", "summary"),
        ("决议事项", "decisions"),
        ("待办事项", "todos"),
        ("原始转录", "transcript"),
    ]

    for heading, key in sections:
        doc.add_heading(heading, level=1)
        doc.add_paragraph("{{ " + key + " }}")
        doc.add_paragraph()

    path = "templates/template.docx"
    doc.save(path)
    print(f"✅ Word 模板已生成：{path}")


# ── 生成 Markdown 模板
def create_md_template():
    content = """# {{ title }}

**日期**：{{ date }}

---

## 会议纪要

{{ summary }}

---

## 决议事项

{{ decisions }}

---

## 待办事项

{{ todos }}

---

## 原始转录

{{ transcript }}
"""
    path = "templates/template.md"
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"✅ Markdown 模板已生成：{path}")


if __name__ == "__main__":
    create_word_template()
    create_md_template()
    print("\n所有模板已生成到 templates/ 目录")
