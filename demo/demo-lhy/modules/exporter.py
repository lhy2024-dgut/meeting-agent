"""
exporter.py - 文档导出模块
支持：Word (.docx)、Markdown (.md)、PDF
使用 docxtpl 填充 Word 模板，reportlab 生成 PDF
"""

import os
from datetime import datetime
from dotenv import load_dotenv

load_dotenv()

STORAGE_BASE = os.getenv("STORAGE_BASE", "./storage")
OUTPUT_DIR   = os.path.join(STORAGE_BASE, "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def _make_filename(meeting_title: str, ext: str) -> str:
    """生成带时间戳的输出文件名"""
    safe_title = "".join(c for c in meeting_title if c.isalnum() or c in "_ -")[:20]
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    return os.path.join(OUTPUT_DIR, f"{safe_title}_{ts}.{ext}")


# ───────────────────────────────────────
# Word 导出（使用 docxtpl 填充模板）
# ───────────────────────────────────────

def export_word(data: dict, template_path: str = None) -> str:
    """
    填充 Word 模板并输出 .docx 文件。
    data 格式：
    {
        "title": "会议标题",
        "date": "2025-01-01",
        "transcript": "转录全文",
        "summary": "会议纪要",
        "todos": "待办事项",
        "decisions": "决议"
    }
    template_path：Word 模板路径（含 {{ 变量 }} 占位符）
    """
    from docxtpl import DocxTemplate

    # 如果没有指定模板，使用默认模板
    if template_path is None:
        template_path = os.path.join("templates", "template.docx")

    if not os.path.exists(template_path):
        # 没有模板就自动创建一个简单的
        _create_default_word_template(template_path)

    doc = DocxTemplate(template_path)

    context = {
        "title":       data.get("title", "会议纪要"),
        "date":        data.get("date", datetime.now().strftime("%Y年%m月%d日")),
        "summary":     data.get("summary", ""),
        "todos":       data.get("todos", ""),
        "decisions":   data.get("decisions", ""),
        "transcript":  data.get("transcript", ""),
    }

    doc.render(context)
    out_path = _make_filename(data.get("title", "meeting"), "docx")
    doc.save(out_path)
    print(f"[导出] Word 文档已保存：{out_path}")
    return out_path


def _create_default_word_template(save_path: str):
    """自动创建一个默认的 Word 模板（含占位符）"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    os.makedirs(os.path.dirname(save_path), exist_ok=True)
    doc = Document()

    # 标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("{{ title }}")
    run.bold = True
    run.font.size = Pt(18)

    # 日期
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run("日期：{{ date }}")

    doc.add_paragraph()

    # 各节
    for heading, key in [
        ("会议纪要", "summary"),
        ("决议事项", "decisions"),
        ("待办事项", "todos"),
        ("原始转录", "transcript"),
    ]:
        h = doc.add_heading(heading, level=1)
        doc.add_paragraph("{{ " + key + " }}")
        doc.add_paragraph()

    doc.save(save_path)
    print(f"[模板] 默认 Word 模板已创建：{save_path}")


# ───────────────────────────────────────
# Markdown 导出
# ───────────────────────────────────────

def export_markdown(data: dict, template_path: str = None) -> str:
    """
    填充 Markdown 模板并输出 .md 文件。
    模板中用 {{ title }} 等占位符。
    """
    if template_path and os.path.exists(template_path):
        with open(template_path, "r", encoding="utf-8") as f:
            content = f.read()
        # 简单字符串替换
        for key, value in data.items():
            content = content.replace("{{ " + key + " }}", str(value))
            content = content.replace("{{" + key + "}}", str(value))
    else:
        # 没有模板，直接生成规范 Markdown
        content = _build_default_markdown(data)

    out_path = _make_filename(data.get("title", "meeting"), "md")
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"[导出] Markdown 文档已保存：{out_path}")
    return out_path


def _build_default_markdown(data: dict) -> str:
    title = data.get("title", "会议纪要")
    date  = data.get("date", datetime.now().strftime("%Y年%m月%d日"))
    return f"""# {title}

**日期**：{date}

---

## 会议纪要

{data.get("summary", "")}

---

## 决议事项

{data.get("decisions", "")}

---

## 待办事项

{data.get("todos", "")}

---

## 原始转录

{data.get("transcript", "")}
"""


# ───────────────────────────────────────
# PDF 导出（用 reportlab 生成，支持中文）
# ───────────────────────────────────────

def export_pdf(data: dict) -> str:
    """
    生成格式规范的 PDF 会议纪要。
    使用 reportlab Platypus 排版引擎，支持自动分页。
    中文字体：使用系统自带的 SimSun/微软雅黑（Windows）
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable, PageBreak
    )
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # 注册中文字体（Windows 系统路径）
    font_paths = [
        "C:/Windows/Fonts/simhei.ttf",    # 黑体
        "C:/Windows/Fonts/simsun.ttc",    # 宋体
        "C:/Windows/Fonts/msyh.ttc",      # 微软雅黑
    ]
    font_name = "SimHei"
    for fp in font_paths:
        if os.path.exists(fp):
            try:
                pdfmetrics.registerFont(TTFont(font_name, fp))
                break
            except Exception:
                continue

    out_path = _make_filename(data.get("title", "meeting"), "pdf")
    doc = SimpleDocTemplate(
        out_path,
        pagesize=A4,
        leftMargin=2.5*cm, rightMargin=2.5*cm,
        topMargin=2.5*cm, bottomMargin=2.5*cm
    )

    styles = getSampleStyleSheet()

    # 自定义样式（中文字体）
    title_style = ParagraphStyle(
        "CnTitle",
        parent=styles["Title"],
        fontName=font_name,
        fontSize=20,
        textColor=colors.HexColor("#1a1a2e"),
        spaceAfter=6,
        alignment=1  # 居中
    )
    h1_style = ParagraphStyle(
        "CnH1",
        parent=styles["Heading1"],
        fontName=font_name,
        fontSize=14,
        textColor=colors.HexColor("#16213e"),
        spaceBefore=14,
        spaceAfter=6,
        borderPadding=(0, 0, 4, 0),
    )
    body_style = ParagraphStyle(
        "CnBody",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=11,
        leading=18,
        spaceAfter=4,
    )
    meta_style = ParagraphStyle(
        "CnMeta",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=10,
        textColor=colors.grey,
        alignment=1
    )

    story = []

    # 标题
    story.append(Paragraph(data.get("title", "会议纪要"), title_style))
    story.append(Paragraph(f"日期：{data.get('date', datetime.now().strftime('%Y年%m月%d日'))}", meta_style))
    story.append(Spacer(1, 0.4*cm))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#cccccc")))
    story.append(Spacer(1, 0.3*cm))

    # 各节内容
    sections = [
        ("会议纪要", data.get("summary", "")),
        ("决议事项", data.get("decisions", "")),
        ("待办事项", data.get("todos", "")),
        ("原始转录", data.get("transcript", "")),
    ]

    for heading, content in sections:
        story.append(Paragraph(heading, h1_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#eeeeee")))
        story.append(Spacer(1, 0.2*cm))

        # 按行分段，处理 bullet points
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                story.append(Spacer(1, 0.15*cm))
                continue
            story.append(Paragraph(line, body_style))

        story.append(Spacer(1, 0.3*cm))

    doc.build(story)
    print(f"[导出] PDF 文档已保存：{out_path}")
    return out_path
